import requests
import json
from DrissionPage import ChromiumPage
from docx import Document
import re
import os

con={
    '6':{
        'id':'6',
        'name':'软件项目管理'
    },
    '6.1':{
        'id':'6.1',
        'name':'项目管理概述',
        'level':'2',
        'tags':'难点',
        'cognitive_dimension':'',
        'teaching_objective':'',
        'content':'content',
    },
    '6.2':{
        'id':'6.2',
        'name':'度量',
        'level':'2',
        'tags':'难点',
        'cognitive_dimension':'',
        'teaching_objective':'',
        'content':'content',
    },
    '6.3':{
        'id':'6.3',
        'name':'估算与项目进度安排',
        'level':'2',
        'tags':'难点',
        'cognitive_dimension':'',
        'teaching_objective':'',
        'content':'content',
    },
    '6.4':{
        'id':'6.4',
        'name':'软件风险管理',
        'level':'2',
        'tags':'',
        'cognitive_dimension':'记忆',
        'teaching_objective':'',
        'content':'content',
    },
    '6.5':{
        'id':'6.5',
        'name':'软件质量属性',
        'level':'2',
        'tags':'难点',
        'cognitive_dimension':'',
        'teaching_objective':'',
        'content':'content',
    },
    '6.6':{
        'id':'6.6',
        'name':'软件质量控制过程',
        'level':'2',
        'tags':'',
        'cognitive_dimension':'记忆',
        'teaching_objective':'',
        'content':'content',
    },
    '6.7':{
        'id':'6.7',
        'name':'软件质量评审',
        'level':'2',
        'tags':'',
        'cognitive_dimension':'记忆',
        'teaching_objective':'',
        'content':'content',
    },
    '6.8':{
        'id':'6.8',
        'name':'软件配置管理',
        'level':'2',
        'tags':'',
        'cognitive_dimension':'记忆',
        'teaching_objective':'',
        'content':'content',
    },
}

def write_knowledge_point_to_word(file_name, case, id_index):
    # 创建一个新的 Document 对象
    doc = Document(file_name)
    if case==2:
        # 添加 "KnowledgePoint" 内容
        doc.add_paragraph('KnowledgePoint')
        # 添加每一项内容
        fields = [
            ('id', id_index),
            ('name', ''),
            ('level', ''),
            ('tags', ''),
            ('cognitive_dimension', ''),
            ('teaching_objective', ''),
            ('content', '')
        ]
        # 逐个添加字段
        for field, value in fields:
            doc.add_paragraph(f'{field}: {value}')
        # 保存文档
        #doc.save(file_name)
    elif case==1:
        # 添加 "Theme" 内容
        doc.add_paragraph('Theme')
        # 添加每一项内容
        fields = [
            ('id', id_index),
            ('name', ''),
            #('level', ''),
            #('tags', ''),
            #('cognitive_dimension', ''),
            #('teaching_objective', ''),
            #('content', '')
        ]
        # 逐个添加字段
        for field, value in fields:
            doc.add_paragraph(f'{field}: {value}')
        # 保存文档
        #doc.save(file_name)
    else:
        # 添加 "Exercise" 内容
        doc.add_paragraph('Exercise')
        # 添加每一项内容
        fields = [
            ('id', id_index),
            #('name', ''),
            #('level', ''),
            #('tags', ''),
            #('cognitive_dimension', ''),
            ('question', ''),
            ('answer', '')
        ]
        # 逐个添加字段
        for field, value in fields:
            doc.add_paragraph(f'{field}: {value}')
        # 保存文档
        #doc.save(file_name)
    doc.add_paragraph()
    doc.save(file_name)

def write_pro(file_name, case, content,user_input):
    # 打开原有的doc对象
    doc = Document(file_name)
    if case==2:
        # 添加 "KnowledgePoint" 内容
        doc.add_paragraph('KnowledgePoint')
        # 添加每一项内容
        fields = [
            ('id', content['id']),  # id对应content中的id
            ('name', content['name']),
            ('level', content['level']),
            ('tags', content['tags']),
            ('cognitive_dimension', content['cognitive_dimension']),
            ('teaching_objective', content['teaching_objective']),
            ('content', user_input)
        ]

        # 逐个添加字段
        for field, value in fields:
            doc.add_paragraph(f'{field}: {value}')
        # 保存文档
        #doc.save(file_name)
    elif case==1:
        # 添加 "Theme" 内容
        doc.add_paragraph('Theme')
        # 添加每一项内容
        fields = [
            ('id', content['id']),
            ('name', content['name']),
            #('level', ''),
            #('tags', ''),
            #('cognitive_dimension', ''),
            #('teaching_objective', ''),
            #('content', '')
        ]
        # 逐个添加字段
        for field, value in fields:
            doc.add_paragraph(f'{field}: {value}')
        # 保存文档
        #doc.save(file_name)
    else:
        # 添加 "Exercise" 内容
        doc.add_paragraph('Exercise')
        # 添加每一项内容
        fields = [
            ('id', content['id']),
            #('name', ''),
            #('level', ''),
            #('tags', ''),
            #('cognitive_dimension', ''),
            ('question', content['question']),
            ('answer', content['answer'])
        ]
        # 逐个添加字段
        for field, value in fields:
            doc.add_paragraph(f'{field}: {value}')
        # 保存文档
        #doc.save(file_name)
    doc.add_paragraph()
    doc.save(file_name)

dialogue1=[]
dialogue1.extend([
    {
        "role": "user",
        "content": "现在你是一名教授软件工程的教师，要根据提供给你的知识点清单，将该内容分类，你只需要提供知识点的编号即可。"
                   "以下是知识点清单："
                   "6软件项目管理			前置知识点：5"	
                   "6.1	项目管理概述			后置知识点：6.2"
                   "6.2	度量		前置知识点：6.1	后置知识点：6.3"
                   "6.3	估算与项目进度安排		前置知识点：6.2	后置知识点：6.4"
                   "6.4	软件风险管理		前置知识点：6.3	后置知识点：6.5"
                   "6.5	软件质量属性		前置知识点：6.4	后置知识点：6.6"
                   "6.6	软件质量控制过程		前置知识点：6.5	后置知识点：6.7"
                   "6.7	软件质量评审		前置知识点：6.6	后置知识点：6.8"
                   "6.8	软件配置管理		前置知识点：6.7"    
                   "我提供一个示例，你接下来的所有回答都参照这个示例回答："
                   "用户：以下是第一段内容：软件项目风险管理是软件项目管理的重要内容。"
                   "你的回答：6.4"
                   "以上是一个示例。"
    },
    {
        "role": "assistant",
        "content": "好的，我明白了。我所有的回答都会参照知识点清单和示例回答，并且绝对不会回答除了知识点编号以外的任何内容。"
    }
])

dialogue2=[]
dialogue2.extend([{
        "role": "user",
        "content": "现在你是一个教授软件工程的教师，你要根据提供给你的知识点内容，将知识点转换成问题，同时也要提供题目的答案。"
                   "我提供一个示例，你接下来的所有回答都参照这个示例回答："
                   "用户：软件项目风险管理是软件项目管理的重要内容。在进行软件项目风险管理时，要辩识风险，评估它们出现的概率及产生的影响，然后建立一个规划来管理风险。风险管理的主要目标是预防风险。"
                   "你的回答：问题：在进行软件项目风险管理时，有哪些要注意的？答案：在进行软件项目风险管理时，要辩识风险，评估它们出现的概率及产生的影响，然后建立一个规划来管理风险。风险管理的主要目标是预防风险。"
                   "除了问题和答案以外，你绝不能回答任何内容。并且问题和答案之间不需要换行。"
},{
        "role": "assistant",
        "content": "好的，我明白了。我所有的回答都会参照知识点内容和示例回答，并且绝对不会回答除了问题和答案以外的任何内容，问题和答案之间也绝不会换行。"
    }])

def get_access_token():
    """
  使用 API Key，Secret Key 获取access_token，替换下列示例中的应用API Key、应用Secret Key
  """

    url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=L0YxfcgdS2PPQ1rRVhORWyni&client_secret=g9g9wDgJKK81fPu61oBISaJYbYRnPR8T"

    payload = json.dumps("")
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
    }

    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json().get("access_token")

def chat(file_name, dialogue):
    user_input = input("用户输入内容(输入“退出”以退出程序): ")  # 获取用户输入
    if user_input == "退出":
        #Document(file_name).add_paragraph("")
        exit(0)
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie-3.5-8k-preview?access_token=" + get_access_token()
    dialogue.append({
                "role": "user",
                "content": user_input
            })
    payload = json.dumps({
        "messages": dialogue
    })
    headers = {
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    # 将响应解析为 JSON，并提取汉字部分
    response_json = response.json()
    #print(response_json['result'])
    answer=response_json.get('result', '')
    #print(answer)
    result = bool(re.match(r"^-?\d+(\.\d+)?$", answer))
    if result==False:
        text=answer
        #text = """问题：……
        #答案：……"""
        # 使用正则表达式进行跨行匹配
        match = re.search(r"问题：(.+?)\n答案：(.+)", text, re.DOTALL)
        if match:
            question = match.group(1).strip()  # 去除多余的空白字符
            answer = match.group(2).strip()
            #print(f"问题: {question}")
            #print(f"答案: {answer}")
            id=quick_determine(user_input)
            cont={
                'id':id,
                'question':question,
                'answer':answer
            }
            write_pro(file_name,3,cont,user_input)
    elif answer=='6':
        write_pro(file_name,1,con[answer],user_input)
    else:
        write_pro(file_name,2,con[answer],user_input)

def quick_determine(user_input)->str:
    flag = False
    answer= ""
    while(flag == False):
        dialogue = []
        dialogue = dialogue1
        url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie-3.5-8k-preview?access_token=" + get_access_token()
        dialogue.append({
            "role": "user",
            "content": user_input
        })
        payload = json.dumps({
            "messages": dialogue
        })
        headers = {
            'Content-Type': 'application/json'
        }
        response = requests.request("POST", url, headers=headers, data=payload)
        response_json = response.json()
        answer = response_json.get('result', '')
        flag = bool(re.match(r"^-?\d+(\.\d+)?$", answer))
    return answer

if __name__ == '__main__':
    file_name = 'file_name.docx'
    if os.path.exists(file_name) == False:
        Document().save(file_name)
    while(True):
      chat('file_name.docx',dialogue2)
      #可修改的chat实参，实参为dialogue2时生成题目，否则生成知识点



